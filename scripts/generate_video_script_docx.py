from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT_PATH = "docs/guion_video_proyecto.docx"


def add_quote(document, text):
    paragraph = document.add_paragraph()
    paragraph.style = "Intense Quote"
    run = paragraph.add_run(text)
    run.italic = True


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def main():
    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = document.add_heading(
        "Guion para Video del Proyecto: Prediccion de Precios de Vivienda en Aguascalientes",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Equipo de 5 integrantes").bold = True

    document.add_paragraph(
        "Duracion sugerida: 5 a 7 minutos. Cada integrante participa aproximadamente 1 minuto y muestra una parte distinta del proyecto."
    )

    document.add_heading("Integrante 1: Presentacion y Problema", level=1)
    document.add_paragraph("Que dice:")
    add_quote(
        document,
        "Hola, somos el equipo [nombre del equipo] y este es nuestro proyecto: Prediccion de Precios de Vivienda en Aguascalientes. "
        "El objetivo es estimar el precio de listado de una vivienda usando variables como superficie, recamaras, banos, estacionamientos, "
        "tipo de propiedad, ubicacion y amenidades cercanas. Este problema se resolvio como un modelo de regresion supervisada, porque "
        "queremos predecir una variable numerica continua: el precio en pesos mexicanos.",
    )
    document.add_paragraph("Que mostrar:")
    add_bullets(
        document,
        [
            "README.md",
            "Titulo del proyecto",
            "Seccion de objetivo",
            "Estructura general del repositorio: README.md, docs/, src/, data/, results/ y models/",
        ],
    )

    document.add_heading("Integrante 2: Datos y Fuentes", level=1)
    document.add_paragraph("Que dice:")
    add_quote(
        document,
        "Para el proyecto se uso un conjunto de datos de viviendas en Aguascalientes. El script intenta consultar fuentes publicas como "
        "Mercado Libre, y si no hay conexion o resultados suficientes, genera un corpus curado y reproducible. El archivo inicial se guarda "
        "en data/raw/properties_raw.csv. Despues de la limpieza, el dataset final queda en data/processed/properties_model.csv con 140 "
        "registros y 27 variables.",
    )
    document.add_paragraph("Que mostrar:")
    add_bullets(
        document,
        [
            "data/raw/properties_raw.csv",
            "data/processed/properties_model.csv",
            "docs/eda_summary.md",
            "Datos clave: 140 registros procesados, 27 variables, precio promedio de $3,161,500 MXN y superficie construida promedio de 132.8 m2.",
            "Tipos de propiedad: casa, casa en condominio y departamento.",
        ],
    )

    document.add_heading("Integrante 3: Limpieza, Variables y Base de Datos", level=1)
    document.add_paragraph("Que dice:")
    add_quote(
        document,
        "En la etapa de limpieza se eliminaron duplicados, registros incompletos, precios no validos y valores atipicos usando IQR. Tambien "
        "se imputaron valores faltantes con la mediana. Ademas, se crearon nuevas variables como precio por metro cuadrado, distancia al centro, "
        "amenidades cercanas, escuelas, hospitales, comercios y si la propiedad era nueva. Finalmente, se genero una base de datos SQLite normalizada.",
    )
    document.add_paragraph("Que mostrar:")
    add_bullets(
        document,
        [
            "src/02_clean_features.py",
            "docs/database_schema.md",
            "data/processed/housing_ags.sqlite",
            "Tablas de la base: properties, locations, amenities, amenity_summary y market_indicators.",
            "Explicar que la normalizacion separa informacion del inmueble, ubicacion y amenidades.",
        ],
    )

    document.add_heading("Integrante 4: Modelos de Machine Learning", level=1)
    document.add_paragraph("Que dice:")
    add_quote(
        document,
        "Para la prediccion se compararon dos modelos. El primero fue Ridge Regression, usado como modelo base lineal e interpretable. "
        "El segundo fue Gradient Boosting, que permite capturar relaciones mas complejas. La variable objetivo fue price_mxn, y se hizo una "
        "particion 80/20 para entrenamiento y prueba, usando una semilla fija para que los resultados fueran reproducibles.",
    )
    document.add_paragraph("Que mostrar:")
    add_bullets(
        document,
        [
            "src/03_train_models.py",
            "Carpeta models/",
            "results/metrics.csv",
            "Modelo 1: Ridge Regression.",
            "Modelo 2: Gradient Boosting.",
            "Metricas usadas: MAE, RMSE y R2.",
            "Explicar que se compararon modelos para saber cual predice mejor el precio.",
        ],
    )

    document.add_heading("Integrante 5: Resultados, Graficas y Conclusion", level=1)
    document.add_paragraph("Que dice:")
    add_quote(
        document,
        "Los resultados muestran que el mejor modelo fue Ridge Regression. Obtuvo un MAE aproximado de 218 mil pesos y un R2 de 0.94. "
        "Esto significa que el modelo explica gran parte de la variacion del precio y que, en promedio, se equivoca por alrededor de 218 mil pesos. "
        "Gradient Boosting obtuvo un MAE de 315 mil pesos y un R2 de 0.87. Aunque es un modelo mas complejo, no supero a Ridge porque el dataset "
        "tiene una relacion principalmente lineal entre precio, superficie y ubicacion.",
    )
    document.add_paragraph("Que mostrar:")
    add_bullets(
        document,
        [
            "results/metrics.csv",
            "results/figures/model_comparison.svg",
            "results/figures/actual_vs_predicted.svg",
            "results/figures/residuals.svg",
            "results/figures/feature_importance.svg",
            "Variables importantes: superficie construida, recamaras, ubicacion, estacionamientos, superficie de terreno y banos.",
        ],
    )
    document.add_paragraph("Cierre sugerido:")
    add_quote(
        document,
        "Como conclusion, el proyecto permite estimar precios de vivienda en Aguascalientes usando datos del inmueble y ubicacion. El modelo "
        "puede servir como estimador inicial o herramienta comparativa, aunque no sustituye una valuacion profesional. Como mejoras futuras se "
        "podrian integrar mas datos reales, ampliar el proyecto a otros estados y probar modelos como Random Forest, XGBoost o LightGBM.",
    )

    document.add_heading("Orden Recomendado Para Grabar", level=1)
    ordered_steps = [
        "Mostrar el repositorio completo.",
        "Abrir README.md.",
        "Mostrar data/raw y data/processed.",
        "Abrir docs/eda_summary.md.",
        "Mostrar src/02_clean_features.py.",
        "Mostrar docs/database_schema.md.",
        "Mostrar src/03_train_models.py.",
        "Mostrar results/metrics.csv.",
        "Mostrar las graficas en results/figures.",
        "Terminar con docs/project_report.md o el README.md.",
    ]
    for step in ordered_steps:
        document.add_paragraph(step, style="List Number")

    document.add_heading("Frase Final Del Equipo", level=1)
    add_quote(
        document,
        "Con esto demostramos un flujo completo de Big Data Analytics: recoleccion de datos, limpieza, ingenieria de variables, base de datos "
        "normalizada, entrenamiento de modelos, evaluacion e interpretacion de resultados.",
    )

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
